import pandas as pd
import streamlit as st
import sympy as sp
import io
from docx import Document


# ---------- TEMPLATE FILLER ----------
def fill_invitation(template_path, data):
    doc = Document(template_path)

    def replace_text_in_paragraph(paragraph, replacements):
        """Replaces placeholders in a paragraph even if split across runs, preserving formatting."""
        full_text = "".join(run.text for run in paragraph.runs)

        # If no placeholder in this paragraph, skip
        if not any(key in full_text for key in replacements.keys()):
            return

        # Replace all placeholders in the full text
        for key, value in replacements.items():
            full_text = full_text.replace(key, value)

        # Clear existing runs
        for run in paragraph.runs:
            run.text = ""

        # Add new run with replaced text (same formatting as first run for safety)
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)

    # Apply to all paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, data)

    # Apply to all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, data)

    # Save into memory
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream


# Initial values of Variables
Pu = None
ρ = None
second_side = None
bar_no = None
nominal_capacity = None

st.header('Design of Short Column (ACI 318-19)',
          divider='rainbow')

# Given data
st.subheader('Data')

with st.container(key='container-1',
                  border=True):

    column_id = st.text_input('Column ID',
                              value=None,
                              placeholder='C1, A1 etc.')

    col1, col2, col3 = st.columns(3, gap='medium', border=True)

    # Loads of Column
    with col1:
        st.write('#### **Loads**')

        dead_load = st.number_input("Dead Load",
                                    value=None,
                                    placeholder='kip')

        live_load = st.number_input("Live Load",
                                    value=None,
                                    placeholder='kip')

        Pu = st.number_input("Ultimate Load (Directly Provided)",
                             value=None,
                             placeholder='kip',
                             help='Provide directly value of ultimate load if you are not providing value of dead and live load')

    # Material Properties of Column
    with col2:
        st.write('#### **Material Properties**')

        f_c = st.number_input("Concrete Strength (f'c)",
                              value=None,
                              placeholder='ksi')

        f_y = st.number_input("Steel Yield Strength (fy)",
                              value=None,
                              placeholder='ksi')
    # Safety Factors
    with col3:
        st.write('#### **Safety Factors**')

        phi = st.number_input("∅",
                              value=0.65)

        k_factor = st.number_input("K",
                                   value=0.8)

# Calculations
st.subheader('Calculations')

with st.container(key='container-2',
                  border=True):
    if dead_load is not None and live_load is not None:
        col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

        # Calculating Ultimate Load
        with col1:
            st.write('Ultimate Load (Pu)')

        with col2:
            Pu = round(1.2*dead_load + 1.6*live_load, 2)
            st.code(f"{Pu} kip")
    else:
        col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

        # Showing Ultimate Load (if directly provided)
        with col1:
            st.write('Ultimate Load (Pu)')

        with col2:
            Pu = Pu
            st.code(f"{Pu} kip")

    if Pu is not None:
        col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

        # Assume Percentage of steel
        with col1:
            st.write('Assume Percentage of Steel (ρ%)')

        with col2:
            ρ = st.number_input('Percentage of Steel (ρ%)',
                                value=None,
                                placeholder='%')

        # Calculate Gross Area (Ag)
    if Pu is not None and ρ is not None and f_c is not None and f_y is not None:
        col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

        with col1:
            st.write('Gross Area (Ag)')

        with col2:
            # Define symbols
            Ag = sp.symbols('Ag')

            # Calculate As from ρ
            ρ = ρ / 100  # Convert percentage to decimal

            # Define the given equation
            equation = sp.Eq(Pu, phi * k_factor * (0.85 * f_c * (Ag - ρ * Ag) + ρ * Ag * f_y))  # As = ρ*Ag

            # Solve for Ag by substituting As
            Ag_calculated = sp.solve(equation, Ag)

            # Display result for gross area (Ag)
            if Ag_calculated:
                st.code(f"{Ag_calculated[0]:.3f} in²")
            else:
                st.error("No solution found for Ag.")

        col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

        # Provide dimension of one side from user
        with col1:
            st.write('Provide Dimension for One Side')

        with col2:
            one_side = st.number_input('Dimension of One Side',
                                       value=None,
                                       placeholder='inch')
        if one_side is not None:
            col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

            # Calculated dimension of second side
            with col1:
                st.write('Dimension of Second Side (calculated)')

            with col2:
                second_side_calculated = Ag_calculated[0] / one_side
                st.code(f'{round(second_side_calculated, 2):.2f} inch')

            col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

            # Provide dimension of second side from user
            with col1:
                st.write('Provide Dimension of Second Side')

            with col2:
                second_side = st.number_input('Dimension of Second Side',
                                              value=None,
                                              placeholder='inch')

        if second_side is not None:
            col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

            # Calculating Ag provided
            with col1:
                st.write('Area of Cross Section - Ag (provided)')

            with col2:
                Ag_provided = one_side * second_side
                st.code(f'{Ag_provided} in²')

            col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

            # Calculating As provided
            with col1:
                st.write('Area of Steel - As (calculated)')

            with col2:
                As_calculated = round(ρ * Ag_provided, 3)
                st.code(f'{As_calculated} in²')

            col1, col2 = st.columns(2, gap='small', vertical_alignment='top')

            # Creating table for no. of bars
            with col1:
                st.write('No of bars (calculated)')

            with col2:
                bar_no_data = pd.DataFrame(data={'Bar No.': [f'#{x}' for x in range(3, 15)],
                                                 'No. of Bars': [As_calculated / ((3.1415 / 4) * (x / 8) ** 2) for x in
                                                                 range(3, 15)]
                                                 })
                bar_no_table = st.dataframe(data=bar_no_data,
                                            hide_index=True,
                                            width='stretch',
                                            height=210)

            col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

            # User Input for finalizing No. of bars
            with col1:
                st.write('Select Bar Number & Enter No. of Bars')

            with col2:
                col2a, col2b = st.columns(2, gap='small', vertical_alignment='top')

                with col2a:
                    bar_no = st.number_input('Bar No.',
                                             value=None,
                                             min_value=0,
                                             placeholder=None)

                with col2b:
                    no_of_bars = st.number_input('No of bars',
                                                 value=None,
                                                 min_value=0,
                                                 placeholder=None)

        if bar_no is not None and no_of_bars is not None:
            col1, col2 = st.columns(2, gap='small', vertical_alignment='center')

            # Calculating As provided
            with col1:
                st.write('Area of Steel - As (provided)')

            with col2:
                As_provided = round(no_of_bars * (3.1415 / 4) * (bar_no / 8) ** 2, 3)
                st.code(f'{As_provided} in²')

            col1, col2 = st.columns(2, gap='small', vertical_alignment='top')

            # Calculating Nominal Capacity of column (∅Pn)
            with col1:
                st.write('Nominal Capacity of Column (∅Pn)')

            with col2:
                nominal_capacity = phi * k_factor * (0.85 * f_c * (Ag_provided - As_provided) + As_provided * f_y)
                st.code(f'{round(nominal_capacity, 2)} kip')

                if nominal_capacity > Pu:
                    st.success("""
                                ∅Pn > Pu,
                                Hence, Design Successful!
                                """)

                else:
                    st.error('Design Unsuccessful!')

            col1, col2 = st.columns(2, gap='small', vertical_alignment='top')

            # Calculating Nominal Capacity of column (∅Pn)
            with col1:
                st.write('Provided Spacing Between Ties')

            with col2:
                tie_spacing_1 = 48 * 3/8
                tie_spacing_2 = 16 * bar_no/8
                tie_spacing_3 = min(one_side, second_side)
                final_tie_spacing = min(tie_spacing_1, tie_spacing_2, tie_spacing_3)
                st.code(f'#3@{final_tie_spacing}"c/c')


# Report-----------------------------------------------------------------------

if nominal_capacity:
    # Data for placeholders in Word document
    formatted_data = {
        '[column_id]': str(column_id),
        '[f_c]': str(f_c),
        '[f_y]': str(f_y),
        '[Pu]': str(Pu),
        '[phi]': str(phi),
        '[k_factor]': str(k_factor),
        '[ρ]': str(ρ),
        '[Ag_calculated]': str(round(Ag_calculated[0], 3)),
        '[one_side]': str(one_side),
        '[second_side_calculated]': str(round(second_side_calculated, 3)),
        '[second_side]': str(second_side),
        '[Ag_provided]': str(round(Ag_provided, 3)),
        '[As_calculated]': str(round(As_calculated, 3)),
        '[bar_no]': str(bar_no),
        '[no_of_bars]': str(no_of_bars),
        '[As_provided]': str(round(As_provided, 3)),
        '[nominal_capacity]': str(round(nominal_capacity, 3)),
        '[tie_spacing_1]': str(round(tie_spacing_1, 0)),
        '[tie_spacing_2]': str(round(tie_spacing_2, 0)),
        '[tie_spacing_3]': str(round(tie_spacing_3, 0)),
        '[final_tie_spacing]': str(round(final_tie_spacing, 0)),
    }

    generate_report_button = st.button('Generate Report')
    if generate_report_button:
        template_path = './Applications/Reinforced_Concrete_Design/Short_Columns/Short_Column_Design_Report_Template.docx'

        # Generate report in memory
        output_file = fill_invitation(template_path, formatted_data)

        # Download button
        st.download_button(
            label="Download Report (Docx)",
            data=output_file,
            file_name=f"Design Report of Column {column_id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.toast('Report has been generated successfully!')